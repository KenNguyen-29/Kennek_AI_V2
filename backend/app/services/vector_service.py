import hashlib
import re
from functools import lru_cache
from pathlib import Path
from typing import Any

import chromadb
import pandas as pd
from docx import Document as DocxDocument
from llama_index.core import Document
from pypdf import PdfReader

CHROMA_PATH = Path("./chroma_db")
COLLECTION_NAME = "user_knowledge"

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".log",
    ".pdf",
    ".docx",
    ".doc",
    ".csv",
    ".xlsx",
    ".xls",
    ".json",
    ".xml",
    ".yaml",
    ".yml",
    ".py",
    ".js",
    ".html",
}

PRINTABLE_RATIO_RE = re.compile(r"[\x09\x0a\x0d\x20-\x7e\u00a0-\uffff]")


@lru_cache(maxsize=1)
def _get_collection() -> Any:
    client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    return client.get_or_create_collection(name=COLLECTION_NAME)


def reset_knowledge_base() -> None:
    client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    try:
        client.delete_collection(COLLECTION_NAME)
    except Exception:
        pass
    _get_collection.cache_clear()
    _get_collection()


def _is_meaningful_text(text: str) -> bool:
    cleaned = text.strip()
    if len(cleaned) < 20:
        return False
    if cleaned.startswith("%PDF"):
        return False
    if "endstream" in cleaned[:500] or "/FlateDecode" in cleaned[:500]:
        return False
    printable = len(PRINTABLE_RATIO_RE.findall(cleaned))
    return (printable / max(len(cleaned), 1)) >= 0.75


def _chunk_text(text: str, *, chunk_size: int = 1800, overlap: int = 200) -> list[str]:
    cleaned = text.strip()
    if not cleaned:
        return []
    if len(cleaned) <= chunk_size:
        return [cleaned]

    chunks: list[str] = []
    start = 0
    while start < len(cleaned):
        end = min(start + chunk_size, len(cleaned))
        chunks.append(cleaned[start:end])
        if end == len(cleaned):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _make_document(text: str, file_path: Path, **extra_metadata: str) -> Document | None:
    if not _is_meaningful_text(text):
        return None
    return Document(
        text=text.strip(),
        metadata={
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_type": file_path.suffix.lower(),
            **extra_metadata,
        },
    )


def _load_pdf(file_path: Path) -> list[Document]:
    reader = PdfReader(str(file_path))
    documents: list[Document] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        document = _make_document(
            text,
            file_path,
            page=str(page_index),
        )
        if document is not None:
            documents.append(document)
    if not documents:
        raise ValueError(
            f"Không trích xuất được chữ từ PDF (có thể là PDF scan/ảnh): {file_path.name}",
        )
    return documents


def _load_docx(file_path: Path) -> list[Document]:
    doc = DocxDocument(str(file_path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    text = "\n".join(paragraphs)
    document = _make_document(text, file_path)
    if document is None:
        raise ValueError(f"File Word trống hoặc không đọc được: {file_path.name}")
    return [document]


def _load_plain_text(file_path: Path) -> list[Document]:
    raw = file_path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1258", "latin-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError(f"Không decode được file text: {file_path.name}")

    document = _make_document(text, file_path)
    if document is None:
        raise ValueError(f"Nội dung file không hợp lệ: {file_path.name}")
    return [document]


def _dataframe_to_text(df: pd.DataFrame, *, source: Path, sheet_name: str) -> str:
    preview = df.fillna("").astype(str)
    markdown_table = preview.to_markdown(index=False)
    return (
        f"Spreadsheet: {source.name}\n"
        f"Sheet: {sheet_name}\n"
        f"Rows: {len(preview)} | Columns: {list(preview.columns)}\n\n"
        f"{markdown_table}"
    )


def _load_spreadsheet(file_path: Path) -> list[Document]:
    workbook = pd.read_excel(file_path, sheet_name=None, dtype=str)
    documents: list[Document] = []
    for sheet_name, frame in workbook.items():
        text = _dataframe_to_text(frame, source=file_path, sheet_name=str(sheet_name))
        document = _make_document(text, file_path, sheet_name=str(sheet_name))
        if document is not None:
            documents.append(document)
    if not documents:
        raise ValueError(f"Không đọc được sheet nào từ Excel: {file_path.name}")
    return documents


def _load_documents_from_paths(file_paths: list[Path]) -> list[Document]:
    documents: list[Document] = []

    for file_path in file_paths:
        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_path.name}")
        if not file_path.is_file():
            raise ValueError(f"File does not exist: {file_path}")

        if suffix == ".pdf":
            documents.extend(_load_pdf(file_path))
        elif suffix == ".docx":
            documents.extend(_load_docx(file_path))
        elif suffix == ".doc":
            raise ValueError(
                f"Định dạng .doc cũ chưa hỗ trợ. Hãy chuyển sang .docx: {file_path.name}",
            )
        elif suffix in {".xlsx", ".xls"}:
            documents.extend(_load_spreadsheet(file_path))
        else:
            documents.extend(_load_plain_text(file_path))

    return documents


def _upsert_documents(documents: list[Document]) -> int:
    ids: list[str] = []
    texts: list[str] = []
    metadatas: list[dict[str, str]] = []

    for document in documents:
        source = str(document.metadata.get("file_path", ""))
        file_name = str(document.metadata.get("file_name") or Path(source).name or "unknown")
        for chunk_index, chunk in enumerate(_chunk_text(document.get_content())):
            if not _is_meaningful_text(chunk):
                continue
            document_id = hashlib.sha256(
                f"{source}\0{chunk_index}\0{chunk}".encode("utf-8"),
            ).hexdigest()
            ids.append(document_id)
            texts.append(chunk)
            metadatas.append(
                {
                    "file_path": source,
                    "file_name": file_name,
                    "chunk_index": str(chunk_index),
                },
            )

    if texts:
        _get_collection().upsert(ids=ids, documents=texts, metadatas=metadatas)

    return len(texts)


def ingest_documents(folder_path: str) -> int:
    source_folder = Path(folder_path)
    if not source_folder.is_dir():
        raise ValueError(f"Document folder does not exist: {folder_path}")

    file_paths = [
        path
        for path in source_folder.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    return int(ingest_files([str(path) for path in file_paths])["chunk_count"])


def ingest_files(file_paths: list[str]) -> dict[str, Any]:
    paths = [Path(path) for path in file_paths]
    if not paths:
        return {"chunk_count": 0, "file_names": []}

    documents = _load_documents_from_paths(paths)
    chunk_count = _upsert_documents(documents)
    return {
        "chunk_count": chunk_count,
        "file_names": [path.name for path in paths],
    }


def query_vector_store(query_str: str) -> list[str]:
    query = query_str.strip()
    if not query:
        return []

    collection = _get_collection()
    result_count = min(5, collection.count())
    if result_count == 0:
        return []

    results = collection.query(
        query_texts=[query],
        n_results=result_count,
        include=["documents", "metadatas"],
    )
    document_groups = results.get("documents") or []
    metadata_groups = results.get("metadatas") or []
    documents = document_groups[0] if document_groups else []
    metadatas = metadata_groups[0] if metadata_groups else []

    contexts: list[str] = []
    for index, document in enumerate(documents):
        if not document or not _is_meaningful_text(document):
            continue
        metadata = metadatas[index] if index < len(metadatas) else {}
        file_name = metadata.get("file_name", "unknown")
        # Show original-looking name if uuid prefix exists.
        display_name = Path(str(file_name)).name
        if "_" in display_name and len(display_name.split("_", 1)[0]) == 32:
            display_name = display_name.split("_", 1)[1]
        contexts.append(f"[Source: {display_name}]\n{document}")
    return contexts
